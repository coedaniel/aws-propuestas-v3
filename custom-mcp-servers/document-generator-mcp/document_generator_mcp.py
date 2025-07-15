#!/usr/bin/env python3
"""
Custom Document Generator MCP Server
Generates Word documents, Excel files, and PDFs for AWS proposals
"""

import asyncio
import json
import logging
import os
import tempfile
from datetime import datetime
from typing import Dict, Any, List, Optional
import base64
import io

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import pandas as pd
from jinja2 import Template
import boto3

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentGeneratorMCP:
    """Custom MCP Server for document generation"""
    
    def __init__(self):
        self.s3_client = boto3.client('s3') if self._has_aws_credentials() else None
        self.temp_dir = tempfile.mkdtemp()
        
    def _has_aws_credentials(self) -> bool:
        """Check if AWS credentials are available"""
        try:
            session = boto3.Session()
            credentials = session.get_credentials()
            return credentials is not None
        except Exception:
            return False
    
    async def generate_word_document(self, content: Dict[str, Any]) -> Dict[str, Any]:
        """Generate a Word document from structured content"""
        try:
            doc = Document()
            
            # Add title
            title = content.get('title', 'AWS Proposal Document')
            title_paragraph = doc.add_heading(title, 0)
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add metadata
            metadata = content.get('metadata', {})
            if metadata:
                doc.add_heading('Document Information', level=1)
                table = doc.add_table(rows=0, cols=2)
                table.style = 'Table Grid'
                
                for key, value in metadata.items():
                    row = table.add_row()
                    row.cells[0].text = str(key).replace('_', ' ').title()
                    row.cells[1].text = str(value)
            
            # Add sections
            sections = content.get('sections', [])
            for section in sections:
                section_title = section.get('title', 'Section')
                doc.add_heading(section_title, level=1)
                
                section_content = section.get('content', '')
                if section_content:
                    doc.add_paragraph(section_content)
                
                # Add subsections
                subsections = section.get('subsections', [])
                for subsection in subsections:
                    subsection_title = subsection.get('title', 'Subsection')
                    doc.add_heading(subsection_title, level=2)
                    
                    subsection_content = subsection.get('content', '')
                    if subsection_content:
                        doc.add_paragraph(subsection_content)
                
                # Add tables if present
                tables = section.get('tables', [])
                for table_data in tables:
                    self._add_table_to_doc(doc, table_data)
                
                # Add lists if present
                lists = section.get('lists', [])
                for list_data in lists:
                    self._add_list_to_doc(doc, list_data)
            
            # Add AWS services section if present
            aws_services = content.get('aws_services', [])
            if aws_services:
                doc.add_heading('AWS Services', level=1)
                for service in aws_services:
                    service_name = service.get('name', 'AWS Service')
                    doc.add_heading(service_name, level=2)
                    
                    description = service.get('description', '')
                    if description:
                        doc.add_paragraph(description)
                    
                    # Add pricing if available
                    pricing = service.get('pricing', {})
                    if pricing:
                        doc.add_heading('Pricing Information', level=3)
                        pricing_table = doc.add_table(rows=1, cols=2)
                        pricing_table.style = 'Table Grid'
                        
                        header_row = pricing_table.rows[0]
                        header_row.cells[0].text = 'Component'
                        header_row.cells[1].text = 'Cost'
                        
                        for component, cost in pricing.items():
                            row = pricing_table.add_row()
                            row.cells[0].text = str(component)
                            row.cells[1].text = str(cost)
            
            # Save document
            filename = f"aws_proposal_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            filepath = os.path.join(self.temp_dir, filename)
            doc.save(filepath)
            
            # Read file as base64 for return
            with open(filepath, 'rb') as f:
                file_content = base64.b64encode(f.read()).decode('utf-8')
            
            return {
                'success': True,
                'filename': filename,
                'filepath': filepath,
                'content_base64': file_content,
                'size_bytes': os.path.getsize(filepath)
            }
            
        except Exception as e:
            logger.error(f"Error generating Word document: {e}")
            return {
                'success': False,
                'error': str(e)
            }
    
    def _add_table_to_doc(self, doc: Document, table_data: Dict[str, Any]):
        """Add a table to the document"""
        headers = table_data.get('headers', [])
        rows = table_data.get('rows', [])
        
        if not headers or not rows:
            return
        
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        
        # Add headers
        header_row = table.rows[0]
        for i, header in enumerate(headers):
            header_row.cells[i].text = str(header)
        
        # Add data rows
        for row_data in rows:
            row = table.add_row()
            for i, cell_data in enumerate(row_data):
                if i < len(row.cells):
                    row.cells[i].text = str(cell_data)
    
    def _add_list_to_doc(self, doc: Document, list_data: Dict[str, Any]):
        """Add a list to the document"""
        list_type = list_data.get('type', 'bullet')  # bullet or numbered
        items = list_data.get('items', [])
        
        for item in items:
            if list_type == 'numbered':
                doc.add_paragraph(str(item), style='List Number')
            else:
                doc.add_paragraph(str(item), style='List Bullet')
    
    async def generate_excel_report(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Generate an Excel report from structured data"""
        try:
            filename = f"aws_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
            filepath = os.path.join(self.temp_dir, filename)
            
            with pd.ExcelWriter(filepath, engine='openpyxl') as writer:
                # Add summary sheet
                summary_data = data.get('summary', {})
                if summary_data:
                    summary_df = pd.DataFrame([summary_data])
                    summary_df.to_excel(writer, sheet_name='Summary', index=False)
                
                # Add AWS services sheet
                services_data = data.get('aws_services', [])
                if services_data:
                    services_df = pd.DataFrame(services_data)
                    services_df.to_excel(writer, sheet_name='AWS Services', index=False)
                
                # Add pricing sheet
                pricing_data = data.get('pricing', [])
                if pricing_data:
                    pricing_df = pd.DataFrame(pricing_data)
                    pricing_df.to_excel(writer, sheet_name='Pricing', index=False)
                
                # Add custom sheets
                custom_sheets = data.get('custom_sheets', {})
                for sheet_name, sheet_data in custom_sheets.items():
                    if isinstance(sheet_data, list) and sheet_data:
                        sheet_df = pd.DataFrame(sheet_data)
                        sheet_df.to_excel(writer, sheet_name=sheet_name, index=False)
            
            # Read file as base64
            with open(filepath, 'rb') as f:
                file_content = base64.b64encode(f.read()).decode('utf-8')
            
            return {
                'success': True,
                'filename': filename,
                'filepath': filepath,
                'content_base64': file_content,
                'size_bytes': os.path.getsize(filepath)
            }
            
        except Exception as e:
            logger.error(f"Error generating Excel report: {e}")
            return {
                'success': False,
                'error': str(e)
            }
    
    async def generate_proposal_template(self, template_type: str, variables: Dict[str, Any]) -> Dict[str, Any]:
        """Generate a proposal document from a template"""
        try:
            templates = {
                'basic': self._get_basic_proposal_template(),
                'detailed': self._get_detailed_proposal_template(),
                'executive': self._get_executive_proposal_template()
            }
            
            template_content = templates.get(template_type, templates['basic'])
            template = Template(template_content)
            
            # Render template with variables
            rendered_content = template.render(**variables)
            
            # Parse rendered content as JSON to create document structure
            try:
                document_structure = json.loads(rendered_content)
            except json.JSONDecodeError:
                # If not JSON, treat as plain text
                document_structure = {
                    'title': f'AWS Proposal - {template_type.title()}',
                    'sections': [
                        {
                            'title': 'Proposal Content',
                            'content': rendered_content
                        }
                    ]
                }
            
            # Generate Word document from structure
            return await self.generate_word_document(document_structure)
            
        except Exception as e:
            logger.error(f"Error generating proposal template: {e}")
            return {
                'success': False,
                'error': str(e)
            }
    
    def _get_basic_proposal_template(self) -> str:
        """Basic proposal template"""
        return '''
{
    "title": "AWS Cloud Solution Proposal",
    "metadata": {
        "client": "{{ client_name }}",
        "date": "{{ date }}",
        "prepared_by": "{{ prepared_by }}",
        "version": "1.0"
    },
    "sections": [
        {
            "title": "Executive Summary",
            "content": "This proposal outlines a comprehensive AWS cloud solution for {{ client_name }}. The proposed architecture leverages {{ primary_services }} to achieve {{ business_objectives }}."
        },
        {
            "title": "Proposed AWS Services",
            "content": "The following AWS services are recommended for this solution:",
            "lists": [
                {
                    "type": "bullet",
                    "items": {{ aws_services_list }}
                }
            ]
        },
        {
            "title": "Implementation Timeline",
            "content": "The project will be implemented in {{ timeline }} phases over {{ duration }}."
        },
        {
            "title": "Investment Summary",
            "content": "The estimated monthly cost for this solution is {{ estimated_cost }}."
        }
    ]
}
'''
    
    def _get_detailed_proposal_template(self) -> str:
        """Detailed proposal template"""
        return '''
{
    "title": "Comprehensive AWS Cloud Architecture Proposal",
    "metadata": {
        "client": "{{ client_name }}",
        "date": "{{ date }}",
        "prepared_by": "{{ prepared_by }}",
        "version": "1.0",
        "project_code": "{{ project_code }}"
    },
    "sections": [
        {
            "title": "Executive Summary",
            "content": "{{ executive_summary }}"
        },
        {
            "title": "Current State Analysis",
            "content": "{{ current_state }}"
        },
        {
            "title": "Proposed Architecture",
            "content": "{{ proposed_architecture }}",
            "subsections": [
                {
                    "title": "Compute Services",
                    "content": "{{ compute_services }}"
                },
                {
                    "title": "Storage Solutions",
                    "content": "{{ storage_solutions }}"
                },
                {
                    "title": "Networking",
                    "content": "{{ networking }}"
                },
                {
                    "title": "Security",
                    "content": "{{ security }}"
                }
            ]
        },
        {
            "title": "Implementation Plan",
            "content": "{{ implementation_plan }}"
        },
        {
            "title": "Cost Analysis",
            "content": "{{ cost_analysis }}"
        },
        {
            "title": "Risk Assessment",
            "content": "{{ risk_assessment }}"
        }
    ]
}
'''
    
    def _get_executive_proposal_template(self) -> str:
        """Executive summary template"""
        return '''
{
    "title": "AWS Cloud Strategy - Executive Brief",
    "metadata": {
        "client": "{{ client_name }}",
        "date": "{{ date }}",
        "prepared_by": "{{ prepared_by }}",
        "executive_sponsor": "{{ executive_sponsor }}"
    },
    "sections": [
        {
            "title": "Business Case",
            "content": "{{ business_case }}"
        },
        {
            "title": "Strategic Benefits",
            "content": "{{ strategic_benefits }}",
            "lists": [
                {
                    "type": "bullet",
                    "items": {{ benefits_list }}
                }
            ]
        },
        {
            "title": "Investment Overview",
            "content": "{{ investment_overview }}"
        },
        {
            "title": "Next Steps",
            "content": "{{ next_steps }}"
        }
    ]
}
'''
    
    async def upload_to_s3(self, filepath: str, bucket: str, key: str) -> Dict[str, Any]:
        """Upload generated document to S3"""
        try:
            if not self.s3_client:
                return {
                    'success': False,
                    'error': 'S3 client not available - AWS credentials not configured'
                }
            
            self.s3_client.upload_file(filepath, bucket, key)
            
            return {
                'success': True,
                'bucket': bucket,
                'key': key,
                's3_url': f's3://{bucket}/{key}'
            }
            
        except Exception as e:
            logger.error(f"Error uploading to S3: {e}")
            return {
                'success': False,
                'error': str(e)
            }
    
    async def list_generated_documents(self) -> Dict[str, Any]:
        """List all generated documents in temp directory"""
        try:
            files = []
            for filename in os.listdir(self.temp_dir):
                filepath = os.path.join(self.temp_dir, filename)
                if os.path.isfile(filepath):
                    files.append({
                        'filename': filename,
                        'size_bytes': os.path.getsize(filepath),
                        'created': datetime.fromtimestamp(os.path.getctime(filepath)).isoformat(),
                        'modified': datetime.fromtimestamp(os.path.getmtime(filepath)).isoformat()
                    })
            
            return {
                'success': True,
                'files': files,
                'total_files': len(files)
            }
            
        except Exception as e:
            logger.error(f"Error listing documents: {e}")
            return {
                'success': False,
                'error': str(e)
            }

# Global instance
document_generator = DocumentGeneratorMCP()

# MCP Tool Functions
async def generate_word_document(content: Dict[str, Any]) -> Dict[str, Any]:
    """MCP Tool: Generate Word document"""
    return await document_generator.generate_word_document(content)

async def generate_excel_report(data: Dict[str, Any]) -> Dict[str, Any]:
    """MCP Tool: Generate Excel report"""
    return await document_generator.generate_excel_report(data)

async def generate_proposal_template(template_type: str, variables: Dict[str, Any]) -> Dict[str, Any]:
    """MCP Tool: Generate proposal from template"""
    return await document_generator.generate_proposal_template(template_type, variables)

async def upload_to_s3(filepath: str, bucket: str, key: str) -> Dict[str, Any]:
    """MCP Tool: Upload document to S3"""
    return await document_generator.upload_to_s3(filepath, bucket, key)

async def list_generated_documents() -> Dict[str, Any]:
    """MCP Tool: List generated documents"""
    return await document_generator.list_generated_documents()

# Available MCP tools
MCP_TOOLS = {
    'generate_word_document': generate_word_document,
    'generate_excel_report': generate_excel_report,
    'generate_proposal_template': generate_proposal_template,
    'upload_to_s3': upload_to_s3,
    'list_generated_documents': list_generated_documents
}
