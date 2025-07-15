# src/generator.py

from docx import Document
import pandas as pd
import os

def generate_all(project_name, description):
    output_path = f"/tmp/output/{project_name}"
    os.makedirs(output_path, exist_ok=True)

    # DOCX
    doc = Document()
    doc.add_heading(f"Project: {project_name}", level=1)
    doc.add_paragraph(description)
    doc.save(f"{output_path}/{project_name}.docx")

    # TXT
    with open(f"{output_path}/{project_name}.txt", "w") as f:
        f.write(f"Project: {project_name}\n\n{description}")

    # CSV
    df = pd.DataFrame([{"project": project_name, "desc": description}])
    df.to_csv(f"{output_path}/{project_name}.csv", index=False)

    return f"Generated files for '{project_name}' in {output_path}: {project_name}.docx, {project_name}.txt, {project_name}.csv"
