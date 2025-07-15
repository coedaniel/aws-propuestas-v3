"""
Document Generator MCP Server
Generates professional Word documents, CSV files, and technical documentation
"""

import asyncio
import json
import logging
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

import pandas as pd
from docx import Document
from docx.shared import Inches
from fastapi import FastAPI, HTTPException
from jinja2 import Template
from mcp.server import Server
from mcp.types import Tool
from pydantic import BaseModel
import uvicorn

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# FastAPI app for HTTP wrapper
app = FastAPI(title="Document Generator MCP Server", version="1.0.0")

# MCP Server
server = Server("document-generator")

class ProjectInfo(BaseModel):
    name: str
    solution_type: str
    selected_services: List[str]
    requirements: Optional[str] = ""
    timeline: Optional[str] = ""
    budget: Optional[str] = ""
    client_name: Optional[str] = ""
    contact_info: Optional[str] = ""

class DocumentRequest(BaseModel):
    project_info: ProjectInfo
    agent_response: Optional[str] = ""

@server.list_tools()
async def list_tools() -> List[Tool]:
    """List available document generation tools"""
    return [
        Tool(
            name="generate_word_document",
            description="Generate a professional Word document for AWS project proposals",
            inputSchema={
                "type": "object",
                "properties": {
                    "project_info": {
                        "type": "object",
                        "properties": {
                            "name": {"type": "string"},
                            "solution_type": {"type": "string"},
                            "selected_services": {"type": "array", "items": {"type": "string"}},
                            "requirements": {"type": "string"},
                            "timeline": {"type": "string"},
                            "budget": {"type": "string"},
                            "client_name": {"type": "string"},
                            "contact_info": {"type": "string"}
                        },
                        "required": ["name", "solution_type", "selected_services"]
                    },
                    "agent_response": {"type": "string"}
                },
                "required": ["project_info"]
            }
        ),
        Tool(
            name="generate_csv_activities",
            description="Generate CSV file with project implementation activities and timeline",
            inputSchema={
                "type": "object",
                "properties": {
                    "project_info": {
                        "type": "object",
                        "properties": {
                            "name": {"type": "string"},
                            "solution_type": {"type": "string"},
                            "selected_services": {"type": "array", "items": {"type": "string"}},
                            "timeline": {"type": "string"}
                        },
                        "required": ["name", "solution_type", "selected_services"]
                    },
                    "agent_response": {"type": "string"}
                },
                "required": ["project_info"]
            }
        ),
        Tool(
            name="generate_complete_proposal",
            description="Generate complete proposal package with Word document and CSV activities",
            inputSchema={
                "type": "object",
                "properties": {
                    "project_info": {
                        "type": "object",
                        "properties": {
                            "name": {"type": "string"},
                            "solution_type": {"type": "string"},
                            "selected_services": {"type": "array", "items": {"type": "string"}},
                            "requirements": {"type": "string"},
                            "timeline": {"type": "string"},
                            "budget": {"type": "string"},
                            "client_name": {"type": "string"},
                            "contact_info": {"type": "string"}
                        },
                        "required": ["name", "solution_type", "selected_services"]
                    },
                    "agent_response": {"type": "string"}
                },
                "required": ["project_info"]
            }
        )
    ]

def create_word_document(project_info: ProjectInfo, agent_response: str = "") -> Dict[str, Any]:
    """Create a professional Word document for the project proposal"""
    
    doc = Document()
    
    # Title
    title = doc.add_heading(f'Propuesta Técnica: {project_info.name}', 0)
    title.alignment = 1  # Center alignment
    
    # Client information
    if project_info.client_name:
        doc.add_heading('Información del Cliente', level=1)
        doc.add_paragraph(f'Cliente: {project_info.client_name}')
        if project_info.contact_info:
            doc.add_paragraph(f'Contacto: {project_info.contact_info}')
    
    # Executive Summary
    doc.add_heading('Resumen Ejecutivo', level=1)
    doc.add_paragraph(
        f'Esta propuesta presenta una solución {project_info.solution_type} '
        f'para el proyecto "{project_info.name}" utilizando servicios de AWS.'
    )
    
    # Project Requirements
    if project_info.requirements:
        doc.add_heading('Requerimientos del Proyecto', level=1)
        doc.add_paragraph(project_info.requirements)
    
    # AWS Services
    doc.add_heading('Servicios AWS Propuestos', level=1)
    for service in project_info.selected_services:
        doc.add_paragraph(f'• {service}', style='List Bullet')
    
    # Technical Solution
    if agent_response:
        doc.add_heading('Solución Técnica Detallada', level=1)
        doc.add_paragraph(agent_response)
    
    # Timeline
    if project_info.timeline:
        doc.add_heading('Cronograma', level=1)
        doc.add_paragraph(project_info.timeline)
    
    # Budget
    if project_info.budget:
        doc.add_heading('Presupuesto Estimado', level=1)
        doc.add_paragraph(project_info.budget)
    
    # Footer
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.add_run('Generado automáticamente por AWS Propuestas V3').italic = True
    footer_para.add_run(f' - {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}').italic = True
    
    # Save document
    filename = f"{project_info.name.replace(' ', '_')}_propuesta_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = Path(f"/tmp/{filename}")
    doc.save(str(filepath))
    
    return {
        "success": True,
        "filename": filename,
        "filepath": str(filepath),
        "message": f"Documento Word generado exitosamente: {filename}"
    }

def create_csv_activities(project_info: ProjectInfo, agent_response: str = "") -> Dict[str, Any]:
    """Create CSV file with project activities and timeline"""
    
    # Base activities for any AWS project
    activities = [
        {
            "Fase": "Planificación",
            "Actividad": "Análisis de requerimientos",
            "Duración (días)": 3,
            "Responsable": "Arquitecto de Soluciones",
            "Dependencias": "Información del cliente"
        },
        {
            "Fase": "Planificación", 
            "Actividad": "Diseño de arquitectura",
            "Duración (días)": 5,
            "Responsable": "Arquitecto de Soluciones",
            "Dependencias": "Análisis de requerimientos"
        },
        {
            "Fase": "Implementación",
            "Actividad": "Configuración de red (VPC, Subnets)",
            "Duración (días)": 2,
            "Responsable": "Ingeniero de Infraestructura",
            "Dependencias": "Diseño de arquitectura"
        }
    ]
    
    # Add service-specific activities
    for service in project_info.selected_services:
        if "EC2" in service:
            activities.append({
                "Fase": "Implementación",
                "Actividad": f"Configuración de instancias {service}",
                "Duración (días)": 3,
                "Responsable": "Ingeniero de Infraestructura",
                "Dependencias": "Configuración de red"
            })
        elif "RDS" in service:
            activities.append({
                "Fase": "Implementación",
                "Actividad": f"Configuración de base de datos {service}",
                "Duración (días)": 2,
                "Responsable": "Administrador de BD",
                "Dependencias": "Configuración de red"
            })
        elif "S3" in service:
            activities.append({
                "Fase": "Implementación",
                "Actividad": f"Configuración de almacenamiento {service}",
                "Duración (días)": 1,
                "Responsable": "Ingeniero de Infraestructura",
                "Dependencias": "Diseño de arquitectura"
            })
        else:
            activities.append({
                "Fase": "Implementación",
                "Actividad": f"Configuración de {service}",
                "Duración (días)": 2,
                "Responsable": "Especialista AWS",
                "Dependencias": "Configuración de red"
            })
    
    # Add final activities
    activities.extend([
        {
            "Fase": "Testing",
            "Actividad": "Pruebas de integración",
            "Duración (días)": 3,
            "Responsable": "QA Engineer",
            "Dependencias": "Implementación completa"
        },
        {
            "Fase": "Despliegue",
            "Actividad": "Despliegue a producción",
            "Duración (días)": 1,
            "Responsable": "DevOps Engineer",
            "Dependencias": "Pruebas exitosas"
        },
        {
            "Fase": "Cierre",
            "Actividad": "Documentación y entrega",
            "Duración (días)": 2,
            "Responsable": "Project Manager",
            "Dependencias": "Despliegue exitoso"
        }
    ])
    
    # Create DataFrame and save to CSV
    df = pd.DataFrame(activities)
    filename = f"{project_info.name.replace(' ', '_')}_actividades_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
    filepath = Path(f"/tmp/{filename}")
    df.to_csv(str(filepath), index=False, encoding='utf-8')
    
    return {
        "success": True,
        "filename": filename,
        "filepath": str(filepath),
        "activities_count": len(activities),
        "message": f"Archivo CSV de actividades generado exitosamente: {filename}"
    }

@server.call_tool()
async def call_tool(name: str, arguments: Dict[str, Any]) -> List[Dict[str, Any]]:
    """Handle tool calls"""
    
    try:
        if name == "generate_word_document":
            project_info = ProjectInfo(**arguments["project_info"])
            agent_response = arguments.get("agent_response", "")
            result = create_word_document(project_info, agent_response)
            return [{"type": "text", "text": json.dumps(result, indent=2)}]
            
        elif name == "generate_csv_activities":
            project_info = ProjectInfo(**arguments["project_info"])
            agent_response = arguments.get("agent_response", "")
            result = create_csv_activities(project_info, agent_response)
            return [{"type": "text", "text": json.dumps(result, indent=2)}]
            
        elif name == "generate_complete_proposal":
            project_info = ProjectInfo(**arguments["project_info"])
            agent_response = arguments.get("agent_response", "")
            
            # Generate both documents
            word_result = create_word_document(project_info, agent_response)
            csv_result = create_csv_activities(project_info, agent_response)
            
            complete_result = {
                "success": True,
                "word_document": word_result,
                "csv_activities": csv_result,
                "message": "Propuesta completa generada exitosamente"
            }
            
            return [{"type": "text", "text": json.dumps(complete_result, indent=2)}]
            
        else:
            raise ValueError(f"Unknown tool: {name}")
            
    except Exception as e:
        logger.error(f"Error in tool {name}: {str(e)}")
        return [{"type": "text", "text": json.dumps({"error": str(e)}, indent=2)}]

# HTTP endpoints for health check and direct API access
@app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {"status": "healthy", "service": "document-generator-mcp"}

@app.post("/generate-word")
async def generate_word_endpoint(request: DocumentRequest):
    """HTTP endpoint for Word document generation"""
    try:
        result = create_word_document(request.project_info, request.agent_response)
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/generate-csv")
async def generate_csv_endpoint(request: DocumentRequest):
    """HTTP endpoint for CSV activities generation"""
    try:
        result = create_csv_activities(request.project_info, request.agent_response)
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/generate-complete")
async def generate_complete_endpoint(request: DocumentRequest):
    """HTTP endpoint for complete proposal generation"""
    try:
        word_result = create_word_document(request.project_info, request.agent_response)
        csv_result = create_csv_activities(request.project_info, request.agent_response)
        
        return {
            "success": True,
            "word_document": word_result,
            "csv_activities": csv_result,
            "message": "Propuesta completa generada exitosamente"
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

async def main():
    """Main function to run both MCP server and HTTP server"""
    # Start HTTP server
    config = uvicorn.Config(app, host="0.0.0.0", port=8005, log_level="info")
    server_instance = uvicorn.Server(config)
    await server_instance.serve()

if __name__ == "__main__":
    asyncio.run(main())
