"""
MCP Server - Professional Document Generation
Generates Word documents and CSV files for AWS proposals
"""

import asyncio
import json
import logging
from datetime import datetime
from typing import Any, Dict, List, Optional
from pathlib import Path

from mcp.server import Server
from mcp.server.models import InitializationOptions
from mcp.server.stdio import stdio_server
from mcp.types import (
    Resource,
    Tool,
    TextContent,
    ImageContent,
    EmbeddedResource,
)
from pydantic import BaseModel

# Document generation imports
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import boto3
import base64
import io

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Initialize MCP server
server = Server("aws-document-generator")

class ProjectInfo(BaseModel):
    name: str
    solution_type: str
    solution_type_detail: Optional[str] = None
    selected_services: List[str] = []
    requirements: Optional[str] = None
    timeline: Optional[str] = None

class DocumentRequest(BaseModel):
    project_info: ProjectInfo
    agent_response: str
    document_type: str  # 'word', 'csv_activities', 'both'

@server.list_tools()
async def handle_list_tools() -> List[Tool]:
    """List available document generation tools"""
    return [
        Tool(
            name="generate_word_document",
            description="Generate professional Word document for AWS proposal",
            inputSchema={
                "type": "object",
                "properties": {
                    "project_info": {
                        "type": "object",
                        "properties": {
                            "name": {"type": "string"},
                            "solution_type": {"type": "string"},
                            "solution_type_detail": {"type": "string"},
                            "selected_services": {"type": "array", "items": {"type": "string"}},
                            "requirements": {"type": "string"},
                            "timeline": {"type": "string"}
                        },
                        "required": ["name", "solution_type"]
                    },
                    "agent_response": {"type": "string"},
                },
                "required": ["project_info", "agent_response"]
            }
        ),
        Tool(
            name="generate_csv_activities",
            description="Generate CSV file with implementation activities timeline",
            inputSchema={
                "type": "object",
                "properties": {
                    "project_info": {
                        "type": "object",
                        "properties": {
                            "name": {"type": "string"},
                            "solution_type": {"type": "string"},
                            "solution_type_detail": {"type": "string"},
                            "selected_services": {"type": "array", "items": {"type": "string"}},
                            "timeline": {"type": "string"}
                        },
                        "required": ["name", "solution_type"]
                    },
                    "agent_response": {"type": "string"},
                },
                "required": ["project_info", "agent_response"]
            }
        ),
        Tool(
            name="generate_complete_proposal",
            description="Generate complete proposal package (Word + CSV)",
            inputSchema={
                "type": "object",
                "properties": {
                    "project_info": {
                        "type": "object",
                        "properties": {
                            "name": {"type": "string"},
                            "solution_type": {"type": "string"},
                            "solution_type_detail": {"type": "string"},
                            "selected_services": {"type": "array", "items": {"type": "string"}},
                            "requirements": {"type": "string"},
                            "timeline": {"type": "string"}
                        },
                        "required": ["name", "solution_type"]
                    },
                    "agent_response": {"type": "string"},
                },
                "required": ["project_info", "agent_response"]
            }
        )
    ]

@server.call_tool()
async def handle_call_tool(name: str, arguments: Dict[str, Any]) -> List[TextContent]:
    """Handle tool calls for document generation"""
    
    try:
        if name == "generate_word_document":
            result = await generate_word_document(arguments)
            return [TextContent(type="text", text=json.dumps(result, indent=2))]
            
        elif name == "generate_csv_activities":
            result = await generate_csv_activities(arguments)
            return [TextContent(type="text", text=json.dumps(result, indent=2))]
            
        elif name == "generate_complete_proposal":
            result = await generate_complete_proposal(arguments)
            return [TextContent(type="text", text=json.dumps(result, indent=2))]
            
        else:
            raise ValueError(f"Unknown tool: {name}")
            
    except Exception as e:
        logger.error(f"Error in tool {name}: {str(e)}")
        return [TextContent(type="text", text=json.dumps({"error": str(e)}))]

async def generate_word_document(arguments: Dict[str, Any]) -> Dict[str, Any]:
    """Generate professional Word document"""
    
    project_info = ProjectInfo(**arguments["project_info"])
    agent_response = arguments["agent_response"]
    
    logger.info(f"Generating Word document for: {project_info.name}")
    
    # Create Word document
    doc = Document()
    
    # Configure document styles
    _configure_document_styles(doc)
    
    # Add document content
    _add_document_header(doc, project_info)
    _add_executive_summary(doc, project_info, agent_response)
    _add_technical_solution(doc, project_info, agent_response)
    _add_implementation_plan(doc, project_info)
    _add_cost_considerations(doc, project_info)
    _add_next_steps(doc, project_info)
    
    # Save document to bytes
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    
    # Encode to base64 for transport
    doc_base64 = base64.b64encode(doc_buffer.getvalue()).decode('utf-8')
    
    filename = f"{project_info.name.replace(' ', '_')}_Propuesta_AWS.docx"
    
    return {
        "filename": filename,
        "content": doc_base64,
        "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "size": len(doc_buffer.getvalue()),
        "generated_at": datetime.now().isoformat()
    }

async def generate_csv_activities(arguments: Dict[str, Any]) -> Dict[str, Any]:
    """Generate CSV with implementation activities"""
    
    project_info = ProjectInfo(**arguments["project_info"])
    agent_response = arguments["agent_response"]
    
    logger.info(f"Generating CSV activities for: {project_info.name}")
    
    # Generate activities based on solution type
    activities = _generate_implementation_activities(project_info)
    
    # Create CSV content
    csv_content = _create_csv_content(activities, project_info)
    
    filename = f"{project_info.name.replace(' ', '_')}_Plan_Actividades.csv"
    
    return {
        "filename": filename,
        "content": csv_content,
        "content_type": "text/csv",
        "size": len(csv_content.encode('utf-8')),
        "generated_at": datetime.now().isoformat(),
        "activities_count": len(activities)
    }

async def generate_complete_proposal(arguments: Dict[str, Any]) -> Dict[str, Any]:
    """Generate complete proposal package"""
    
    word_doc = await generate_word_document(arguments)
    csv_activities = await generate_csv_activities(arguments)
    
    return {
        "word_document": word_doc,
        "csv_activities": csv_activities,
        "package_generated_at": datetime.now().isoformat()
    }

def _configure_document_styles(doc: Document):
    """Configure document styles"""
    
    # Title style
    title_style = doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_font = title_style.font
    title_font.name = 'Calibri'
    title_font.size = Pt(24)
    title_font.bold = True
    title_font.color.rgb = None  # Default color
    
    # Heading styles
    heading1_style = doc.styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
    heading1_font = heading1_style.font
    heading1_font.name = 'Calibri'
    heading1_font.size = Pt(18)
    heading1_font.bold = True
    
    # Normal text style
    normal_style = doc.styles.add_style('CustomNormal', WD_STYLE_TYPE.PARAGRAPH)
    normal_font = normal_style.font
    normal_font.name = 'Calibri'
    normal_font.size = Pt(11)

def _add_document_header(doc: Document, project_info: ProjectInfo):
    """Add document header with title and project info"""
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(f"PROPUESTA TÉCNICA AWS\n{project_info.name.upper()}")
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(24)
    title_run.bold = True
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(f"Solución: {project_info.solution_type.replace('_', ' ').title()}")
    subtitle_run.font.name = 'Calibri'
    subtitle_run.font.size = Pt(16)
    subtitle_run.italic = True
    
    # Date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"Fecha: {datetime.now().strftime('%d de %B de %Y')}")
    date_run.font.name = 'Calibri'
    date_run.font.size = Pt(12)
    
    # Add page break
    doc.add_page_break()

def _add_executive_summary(doc: Document, project_info: ProjectInfo, agent_response: str):
    """Add executive summary section"""
    
    heading = doc.add_heading('RESUMEN EJECUTIVO', level=1)
    heading.style.font.name = 'Calibri'
    
    # Project overview
    overview = doc.add_paragraph()
    overview_text = f"""
Este documento presenta la propuesta técnica para la implementación de "{project_info.name}" 
utilizando servicios de Amazon Web Services (AWS). La solución propuesta está diseñada para 
proporcionar una arquitectura escalable, segura y costo-efectiva que satisfaga los 
requerimientos específicos del proyecto.

Tipo de Solución: {project_info.solution_type.replace('_', ' ').title()}
"""
    
    if project_info.solution_type_detail:
        overview_text += f"Detalle: {project_info.solution_type_detail}\n"
    
    if project_info.selected_services:
        overview_text += f"\nServicios AWS Principales: {', '.join(project_info.selected_services)}"
    
    overview.add_run(overview_text.strip())
    overview.style.font.name = 'Calibri'
    overview.style.font.size = Pt(11)

def _add_technical_solution(doc: Document, project_info: ProjectInfo, agent_response: str):
    """Add technical solution section"""
    
    heading = doc.add_heading('SOLUCIÓN TÉCNICA', level=1)
    heading.style.font.name = 'Calibri'
    
    # Architecture overview
    arch_heading = doc.add_heading('Arquitectura Propuesta', level=2)
    arch_heading.style.font.name = 'Calibri'
    
    # Extract relevant technical details from agent response
    tech_content = _extract_technical_content(agent_response, project_info)
    
    tech_para = doc.add_paragraph(tech_content)
    tech_para.style.font.name = 'Calibri'
    tech_para.style.font.size = Pt(11)
    
    # Services breakdown
    if project_info.selected_services:
        services_heading = doc.add_heading('Servicios AWS Incluidos', level=2)
        services_heading.style.font.name = 'Calibri'
        
        for service in project_info.selected_services:
            service_desc = _get_service_description(service)
            service_para = doc.add_paragraph(f"• {service}: {service_desc}")
            service_para.style.font.name = 'Calibri'
            service_para.style.font.size = Pt(11)

def _add_implementation_plan(doc: Document, project_info: ProjectInfo):
    """Add implementation plan section"""
    
    heading = doc.add_heading('PLAN DE IMPLEMENTACIÓN', level=1)
    heading.style.font.name = 'Calibri'
    
    phases = _get_implementation_phases(project_info)
    
    for phase_num, phase in enumerate(phases, 1):
        phase_heading = doc.add_heading(f'Fase {phase_num}: {phase["name"]}', level=2)
        phase_heading.style.font.name = 'Calibri'
        
        # Duration
        duration_para = doc.add_paragraph(f"Duración estimada: {phase['duration']}")
        duration_para.style.font.name = 'Calibri'
        duration_para.style.font.size = Pt(11)
        duration_para.style.font.bold = True
        
        # Activities
        activities_para = doc.add_paragraph("Actividades principales:")
        activities_para.style.font.name = 'Calibri'
        activities_para.style.font.size = Pt(11)
        activities_para.style.font.bold = True
        
        for activity in phase['activities']:
            activity_para = doc.add_paragraph(f"• {activity}")
            activity_para.style.font.name = 'Calibri'
            activity_para.style.font.size = Pt(11)

def _add_cost_considerations(doc: Document, project_info: ProjectInfo):
    """Add cost considerations section"""
    
    heading = doc.add_heading('CONSIDERACIONES DE COSTOS', level=1)
    heading.style.font.name = 'Calibri'
    
    cost_text = """
Los costos de AWS se basan en un modelo de pago por uso, lo que permite optimizar 
los gastos según las necesidades reales del proyecto. Los principales factores 
que influyen en los costos incluyen:

• Recursos de cómputo (EC2, Lambda)
• Almacenamiento (S3, EBS)
• Transferencia de datos
• Servicios administrados (RDS, ELB)

Se recomienda utilizar la Calculadora de Precios de AWS para obtener estimaciones 
precisas basadas en los patrones de uso específicos del proyecto.
"""
    
    cost_para = doc.add_paragraph(cost_text.strip())
    cost_para.style.font.name = 'Calibri'
    cost_para.style.font.size = Pt(11)

def _add_next_steps(doc: Document, project_info: ProjectInfo):
    """Add next steps section"""
    
    heading = doc.add_heading('PRÓXIMOS PASOS', level=1)
    heading.style.font.name = 'Calibri'
    
    next_steps = [
        "Revisión y aprobación de la propuesta técnica",
        "Definición detallada de requerimientos específicos",
        "Configuración del entorno de desarrollo/pruebas",
        "Inicio de la implementación según el cronograma establecido",
        "Configuración de monitoreo y alertas",
        "Capacitación del equipo técnico",
        "Documentación y transferencia de conocimiento"
    ]
    
    for step in next_steps:
        step_para = doc.add_paragraph(f"• {step}")
        step_para.style.font.name = 'Calibri'
        step_para.style.font.size = Pt(11)

def _extract_technical_content(agent_response: str, project_info: ProjectInfo) -> str:
    """Extract and format technical content from agent response"""
    
    # Basic technical description based on solution type
    if project_info.solution_type == 'rapid_service':
        return f"""
La solución propuesta implementa una arquitectura basada en los servicios AWS seleccionados: 
{', '.join(project_info.selected_services)}. Esta configuración proporciona una base sólida 
para el despliegue rápido de servicios en la nube, con alta disponibilidad y escalabilidad 
automática según las necesidades del negocio.

La arquitectura incluye componentes de red seguros, almacenamiento redundante y capacidades 
de monitoreo integradas para garantizar el rendimiento óptimo del sistema.
"""
    else:
        return f"""
La solución integral propuesta para "{project_info.name}" implementa una arquitectura 
empresarial completa que abarca todos los aspectos necesarios para una implementación 
robusta en AWS. 

{project_info.solution_type_detail or 'La solución está diseñada para proporcionar escalabilidad, seguridad y alta disponibilidad.'}

La arquitectura sigue las mejores prácticas del AWS Well-Architected Framework, 
garantizando seguridad, confiabilidad, eficiencia de rendimiento, optimización 
de costos y excelencia operacional.
"""

def _get_service_description(service: str) -> str:
    """Get description for AWS service"""
    
    descriptions = {
        'EC2': 'Instancias de cómputo escalables para aplicaciones',
        'RDS': 'Base de datos relacional administrada con alta disponibilidad',
        'S3': 'Almacenamiento de objetos escalable y duradero',
        'ELB': 'Balanceador de carga para distribución de tráfico',
        'CloudFront': 'Red de distribución de contenido global',
        'Lambda': 'Cómputo serverless para ejecución de código',
        'VPC': 'Red privada virtual para aislamiento de recursos',
        'IAM': 'Gestión de identidades y accesos',
        'CloudWatch': 'Monitoreo y observabilidad de recursos'
    }
    
    return descriptions.get(service, 'Servicio AWS especializado para la solución')

def _get_implementation_phases(project_info: ProjectInfo) -> List[Dict[str, Any]]:
    """Get implementation phases based on solution type"""
    
    if project_info.solution_type == 'rapid_service':
        return [
            {
                'name': 'Configuración Inicial',
                'duration': '1-2 semanas',
                'activities': [
                    'Configuración de cuenta AWS y permisos',
                    'Creación de VPC y componentes de red',
                    'Configuración de grupos de seguridad'
                ]
            },
            {
                'name': 'Despliegue de Servicios',
                'duration': '1-2 semanas',
                'activities': [
                    'Implementación de servicios seleccionados',
                    'Configuración de balanceadores de carga',
                    'Configuración de almacenamiento'
                ]
            },
            {
                'name': 'Pruebas y Optimización',
                'duration': '1 semana',
                'activities': [
                    'Pruebas de funcionalidad',
                    'Optimización de rendimiento',
                    'Configuración de monitoreo'
                ]
            }
        ]
    else:
        return [
            {
                'name': 'Planificación y Diseño',
                'duration': '2-3 semanas',
                'activities': [
                    'Análisis detallado de requerimientos',
                    'Diseño de arquitectura completa',
                    'Definición de estándares y políticas'
                ]
            },
            {
                'name': 'Configuración de Infraestructura',
                'duration': '3-4 semanas',
                'activities': [
                    'Implementación de red y seguridad',
                    'Configuración de servicios base',
                    'Implementación de alta disponibilidad'
                ]
            },
            {
                'name': 'Desarrollo e Integración',
                'duration': '4-6 semanas',
                'activities': [
                    'Desarrollo de componentes específicos',
                    'Integración de servicios',
                    'Configuración de pipelines CI/CD'
                ]
            },
            {
                'name': 'Pruebas y Despliegue',
                'duration': '2-3 semanas',
                'activities': [
                    'Pruebas integrales del sistema',
                    'Despliegue en producción',
                    'Configuración de monitoreo y alertas'
                ]
            }
        ]

def _generate_implementation_activities(project_info: ProjectInfo) -> List[Dict[str, Any]]:
    """Generate detailed implementation activities"""
    
    activities = []
    
    # Base activities for all projects
    base_activities = [
        {
            'fase': 'Planificación',
            'actividad': 'Análisis de requerimientos técnicos',
            'responsable': 'Arquitecto de Soluciones',
            'duracion_dias': 3,
            'dependencias': '',
            'entregables': 'Documento de requerimientos técnicos'
        },
        {
            'fase': 'Planificación',
            'actividad': 'Diseño de arquitectura AWS',
            'responsable': 'Arquitecto de Soluciones',
            'duracion_dias': 5,
            'dependencias': 'Análisis de requerimientos',
            'entregables': 'Diagrama de arquitectura, Documento técnico'
        },
        {
            'fase': 'Configuración',
            'actividad': 'Configuración de cuenta AWS y permisos',
            'responsable': 'DevOps Engineer',
            'duracion_dias': 2,
            'dependencias': 'Diseño de arquitectura',
            'entregables': 'Cuenta AWS configurada, Políticas IAM'
        },
        {
            'fase': 'Configuración',
            'actividad': 'Creación de VPC y componentes de red',
            'responsable': 'DevOps Engineer',
            'duracion_dias': 3,
            'dependencias': 'Configuración de cuenta',
            'entregables': 'VPC, Subnets, Security Groups'
        }
    ]
    
    activities.extend(base_activities)
    
    # Add service-specific activities
    if 'EC2' in project_info.selected_services:
        activities.extend([
            {
                'fase': 'Implementación',
                'actividad': 'Configuración de instancias EC2',
                'responsable': 'DevOps Engineer',
                'duracion_dias': 3,
                'dependencias': 'VPC configurada',
                'entregables': 'Instancias EC2 operativas'
            }
        ])
    
    if 'RDS' in project_info.selected_services:
        activities.extend([
            {
                'fase': 'Implementación',
                'actividad': 'Configuración de base de datos RDS',
                'responsable': 'Database Administrator',
                'duracion_dias': 4,
                'dependencias': 'VPC configurada',
                'entregables': 'Base de datos RDS operativa'
            }
        ])
    
    if 'S3' in project_info.selected_services:
        activities.extend([
            {
                'fase': 'Implementación',
                'actividad': 'Configuración de buckets S3',
                'responsable': 'DevOps Engineer',
                'duracion_dias': 2,
                'dependencias': 'Configuración de cuenta',
                'entregables': 'Buckets S3 con políticas de seguridad'
            }
        ])
    
    # Final activities
    final_activities = [
        {
            'fase': 'Pruebas',
            'actividad': 'Pruebas de integración',
            'responsable': 'QA Engineer',
            'duracion_dias': 5,
            'dependencias': 'Todos los servicios implementados',
            'entregables': 'Reporte de pruebas'
        },
        {
            'fase': 'Despliegue',
            'actividad': 'Configuración de monitoreo',
            'responsable': 'DevOps Engineer',
            'duracion_dias': 3,
            'dependencias': 'Pruebas completadas',
            'entregables': 'Dashboard de monitoreo, Alertas configuradas'
        },
        {
            'fase': 'Cierre',
            'actividad': 'Documentación y transferencia',
            'responsable': 'Arquitecto de Soluciones',
            'duracion_dias': 3,
            'dependencias': 'Sistema en producción',
            'entregables': 'Documentación técnica, Manual de operación'
        }
    ]
    
    activities.extend(final_activities)
    
    return activities

def _create_csv_content(activities: List[Dict[str, Any]], project_info: ProjectInfo) -> str:
    """Create CSV content from activities"""
    
    csv_lines = [
        f"Plan de Actividades - {project_info.name}",
        f"Generado el: {datetime.now().strftime('%d/%m/%Y %H:%M')}",
        "",
        "Fase,Actividad,Responsable,Duración (días),Dependencias,Entregables"
    ]
    
    for activity in activities:
        line = f'"{activity["fase"]}","{activity["actividad"]}","{activity["responsable"]}",{activity["duracion_dias"]},"{activity["dependencias"]}","{activity["entregables"]}"'
        csv_lines.append(line)
    
    # Add summary
    total_days = sum(activity["duracion_dias"] for activity in activities)
    csv_lines.extend([
        "",
        f"RESUMEN DEL PROYECTO",
        f"Total de actividades: {len(activities)}",
        f"Duración total estimada: {total_days} días",
        f"Duración en semanas: {total_days // 5} semanas",
        "",
        "NOTAS:",
        "- Las duraciones son estimadas y pueden variar según la complejidad",
        "- Algunas actividades pueden ejecutarse en paralelo",
        "- Se recomienda incluir tiempo adicional para contingencias"
    ])
    
    return '\n'.join(csv_lines)

async def main():
    """Main function to run the MCP server"""
    async with stdio_server() as (read_stream, write_stream):
        await server.run(
            read_stream,
            write_stream,
            InitializationOptions(
                server_name="aws-document-generator",
                server_version="1.0.0",
                capabilities=server.get_capabilities(
                    notification_options=None,
                    experimental_capabilities=None,
                ),
            ),
        )

if __name__ == "__main__":
    asyncio.run(main())
