"""
Word document generator for AWS proposals
"""
import io
from typing import Dict, Any
from docx import Document
from docx.shared import Inches

def generate_word_document(project_info: Dict[str, Any], document_type: str = "propuesta") -> bytes:
    """
    Generate a Word document for the project
    
    Args:
        project_info: Dictionary with project information
        document_type: Type of document (propuesta, tecnico, resumen)
    
    Returns:
        bytes: Word document content
    """
    doc = Document()
    
    # Get project details
    project_name = project_info.get('name', 'Proyecto AWS')
    project_type = project_info.get('type', 'general')
    service_type = project_info.get('service_type', 'general')
    description = project_info.get('description', f'Solución {service_type} en AWS')
    objective = project_info.get('objective', 'Implementar una solución robusta y escalable en AWS')
    
    # Title
    title = doc.add_heading(f'Propuesta Ejecutiva - {project_name}', 0)
    
    # Executive Summary
    doc.add_heading('Resumen Ejecutivo', level=1)
    doc.add_paragraph(
        f'Este documento presenta la propuesta técnica y comercial para el proyecto '
        f'"{project_name}", {description}. La solución propuesta cumple con '
        f'los requerimientos especificados y sigue las mejores prácticas de AWS.'
    )
    
    # Project Objective
    doc.add_heading('Objetivo del Proyecto', level=1)
    doc.add_paragraph(objective)
    
    # Project Description
    doc.add_heading('Descripción del Proyecto', level=1)
    doc.add_paragraph(description)
    
    # Project Type and Service Details
    doc.add_heading('Tipo de Proyecto', level=1)
    if project_type == 'servicio_rapido':
        doc.add_paragraph('Este es un proyecto de servicio rápido, enfocado en la implementación específica de servicios AWS.')
    elif project_type == 'solucion_integral':
        doc.add_paragraph('Este es un proyecto de solución integral que abarca múltiples servicios y componentes de AWS.')
    
    # Service-specific details - COMPREHENSIVE SUPPORT FOR ALL SERVICES
    if service_type == 'vpc':
        doc.add_heading('Configuración de VPC', level=1)
        doc.add_paragraph('La solución incluye la configuración de una Virtual Private Cloud (VPC) con las siguientes características:')
        
        vpc_details = doc.add_paragraph()
        vpc_details.add_run('• ').bold = True
        vpc_details.add_run('VPC principal con CIDR block configurable\n')
        vpc_details.add_run('• ').bold = True
        vpc_details.add_run('Subredes públicas y privadas en múltiples zonas de disponibilidad\n')
        vpc_details.add_run('• ').bold = True
        vpc_details.add_run('Internet Gateway para conectividad externa\n')
        vpc_details.add_run('• ').bold = True
        vpc_details.add_run('Tablas de rutas optimizadas\n')
        
        if project_info.get('architecture') == 'tres_capas':
            doc.add_paragraph('La arquitectura está diseñada para soportar un sistema de tres capas (presentación, lógica de negocio y datos).')
        
        # CIDR blocks if specified
        cidr_blocks = project_info.get('cidr_blocks')
        if cidr_blocks:
            doc.add_heading('Configuración de Red', level=2)
            doc.add_paragraph(f'VPC CIDR: {cidr_blocks.get("vpc", "10.0.0.0/16")}')
            if cidr_blocks.get('public_subnets'):
                doc.add_paragraph(f'Subredes Públicas: {", ".join(cidr_blocks["public_subnets"])}')
            if cidr_blocks.get('private_subnets'):
                doc.add_paragraph(f'Subredes Privadas: {", ".join(cidr_blocks["private_subnets"])}')
    
    elif service_type == 'ec2':
        doc.add_heading('Configuración de EC2', level=1)
        doc.add_paragraph('La solución incluye la configuración de instancias Amazon EC2 con las siguientes especificaciones:')
        
        ec2_details = doc.add_paragraph()
        instance_type = project_info.get('instance_type', 't2.micro')
        ec2_details.add_run('• ').bold = True
        ec2_details.add_run(f'Tipo de instancia: {instance_type}\n')
        ec2_details.add_run('• ').bold = True
        ec2_details.add_run('Sistema operativo optimizado\n')
        ec2_details.add_run('• ').bold = True
        ec2_details.add_run('Configuración de seguridad avanzada\n')
        ec2_details.add_run('• ').bold = True
        ec2_details.add_run('Monitoreo y alertas integradas\n')
        
        # Technical specs if available
        tech_specs = project_info.get('technical_specs', {})
        if tech_specs:
            doc.add_heading('Especificaciones Técnicas', level=2)
            for spec, value in tech_specs.items():
                doc.add_paragraph(f'{spec.title()}: {value}')
    
    elif service_type == 'rds':
        doc.add_heading('Configuración de RDS', level=1)
        doc.add_paragraph('La solución incluye una base de datos Amazon RDS con las siguientes características:')
        
        rds_details = doc.add_paragraph()
        rds_details.add_run('• ').bold = True
        rds_details.add_run('Base de datos relacional administrada\n')
        rds_details.add_run('• ').bold = True
        rds_details.add_run('Respaldos automáticos\n')
        rds_details.add_run('• ').bold = True
        rds_details.add_run('Alta disponibilidad con Multi-AZ\n')
        rds_details.add_run('• ').bold = True
        rds_details.add_run('Cifrado en reposo y en tránsito\n')
    
    elif service_type == 's3':
        doc.add_heading('Configuración de Amazon S3', level=1)
        doc.add_paragraph('La solución incluye la configuración de buckets Amazon S3 con las siguientes características:')
        
        s3_details = doc.add_paragraph()
        bucket_name = project_info.get('bucket_name', f'{project_name.lower().replace(" ", "-")}-bucket')
        s3_details.add_run('• ').bold = True
        s3_details.add_run(f'Bucket principal: {bucket_name}\n')
        
        storage_type = project_info.get('storage_type', 'Standard')
        s3_details.add_run('• ').bold = True
        s3_details.add_run(f'Clase de almacenamiento: {storage_type}\n')
        
        aws_region = project_info.get('aws_region', project_info.get('region', 'us-east-1'))
        s3_details.add_run('• ').bold = True
        s3_details.add_run(f'Región: {aws_region}\n')
        
        versioning = project_info.get('versioning', 'enabled')
        s3_details.add_run('• ').bold = True
        s3_details.add_run(f'Versionado: {versioning}\n')
        
        encryption = project_info.get('encryption', 'enabled')
        s3_details.add_run('• ').bold = True
        s3_details.add_run(f'Cifrado: {encryption}\n')
        
        access_policy = project_info.get('access_policy', 'private')
        s3_details.add_run('• ').bold = True
        s3_details.add_run(f'Política de acceso: {access_policy}\n')
    
    elif service_type == 'efs':
        doc.add_heading('Configuración de Amazon EFS', level=1)
        doc.add_paragraph('La solución incluye la configuración de Amazon Elastic File System (EFS) con las siguientes características:')
        
        efs_details = doc.add_paragraph()
        efs_details.add_run('• ').bold = True
        efs_details.add_run('Sistema de archivos elástico y escalable\n')
        efs_details.add_run('• ').bold = True
        efs_details.add_run('Acceso concurrente desde múltiples instancias EC2\n')
        efs_details.add_run('• ').bold = True
        efs_details.add_run('Cifrado en reposo y en tránsito\n')
        efs_details.add_run('• ').bold = True
        efs_details.add_run('Respaldos automáticos configurados\n')
        efs_details.add_run('• ').bold = True
        efs_details.add_run('Políticas de ciclo de vida para optimización de costos\n')
        
        performance_mode = project_info.get('performance_mode', 'General Purpose')
        efs_details.add_run('• ').bold = True
        efs_details.add_run(f'Modo de rendimiento: {performance_mode}\n')
        
        throughput_mode = project_info.get('throughput_mode', 'Bursting')
        efs_details.add_run('• ').bold = True
        efs_details.add_run(f'Modo de throughput: {throughput_mode}\n')
    
    elif service_type == 'lambda':
        doc.add_heading('Configuración de AWS Lambda', level=1)
        doc.add_paragraph('La solución incluye funciones AWS Lambda con las siguientes características:')
        
        lambda_details = doc.add_paragraph()
        runtime = project_info.get('runtime', 'python3.9')
        lambda_details.add_run('• ').bold = True
        lambda_details.add_run(f'Runtime: {runtime}\n')
        lambda_details.add_run('• ').bold = True
        lambda_details.add_run('Ejecución serverless sin gestión de infraestructura\n')
        lambda_details.add_run('• ').bold = True
        lambda_details.add_run('Escalado automático basado en demanda\n')
        lambda_details.add_run('• ').bold = True
        lambda_details.add_run('Integración con otros servicios AWS\n')
        lambda_details.add_run('• ').bold = True
        lambda_details.add_run('Monitoreo con CloudWatch integrado\n')
    
    elif service_type == 'cloudfront':
        doc.add_heading('Configuración de Amazon CloudFront', level=1)
        doc.add_paragraph('La solución incluye una distribución CloudFront con las siguientes características:')
        
        cf_details = doc.add_paragraph()
        cf_details.add_run('• ').bold = True
        cf_details.add_run('Red de distribución de contenido global\n')
        cf_details.add_run('• ').bold = True
        cf_details.add_run('Reducción de latencia para usuarios finales\n')
        cf_details.add_run('• ').bold = True
        cf_details.add_run('Certificado SSL/TLS incluido\n')
        cf_details.add_run('• ').bold = True
        cf_details.add_run('Compresión automática de contenido\n')
        cf_details.add_run('• ').bold = True
        cf_details.add_run('Protección DDoS integrada\n')
    
    elif service_type == 'elb':
        doc.add_heading('Configuración de Elastic Load Balancer', level=1)
        doc.add_paragraph('La solución incluye un balanceador de carga con las siguientes características:')
        
        elb_details = doc.add_paragraph()
        elb_details.add_run('• ').bold = True
        elb_details.add_run('Distribución automática de tráfico\n')
        elb_details.add_run('• ').bold = True
        elb_details.add_run('Health checks automáticos\n')
        elb_details.add_run('• ').bold = True
        elb_details.add_run('Alta disponibilidad multi-AZ\n')
        elb_details.add_run('• ').bold = True
        elb_details.add_run('Terminación SSL/TLS\n')
        elb_details.add_run('• ').bold = True
        elb_details.add_run('Integración con Auto Scaling\n')
    
    elif service_type == 'ses':
        doc.add_heading('Configuración de Amazon SES', level=1)
        doc.add_paragraph('La solución incluye Amazon Simple Email Service con las siguientes características:')
        
        ses_details = doc.add_paragraph()
        ses_details.add_run('• ').bold = True
        ses_details.add_run('Envío de correos electrónicos escalable\n')
        ses_details.add_run('• ').bold = True
        ses_details.add_run('Verificación de dominios y direcciones\n')
        ses_details.add_run('• ').bold = True
        ses_details.add_run('Métricas de entrega y rebotes\n')
        ses_details.add_run('• ').bold = True
        ses_details.add_run('Configuración de DKIM y SPF\n')
        ses_details.add_run('• ').bold = True
        ses_details.add_run('Gestión de listas de supresión\n')
    
    elif service_type == 'vpn':
        doc.add_heading('Configuración de AWS VPN', level=1)
        doc.add_paragraph('La solución incluye una conexión VPN con las siguientes características:')
        
        vpn_details = doc.add_paragraph()
        vpn_details.add_run('• ').bold = True
        vpn_details.add_run('Conexión segura site-to-site\n')
        vpn_details.add_run('• ').bold = True
        vpn_details.add_run('Cifrado IPSec estándar de la industria\n')
        vpn_details.add_run('• ').bold = True
        vpn_details.add_run('Redundancia con múltiples túneles\n')
        vpn_details.add_run('• ').bold = True
        vpn_details.add_run('Monitoreo de conectividad\n')
        vpn_details.add_run('• ').bold = True
        vpn_details.add_run('Configuración de rutas automática\n')
    
    elif service_type == 'backup':
        doc.add_heading('Configuración de AWS Backup', level=1)
        doc.add_paragraph('La solución incluye un sistema de respaldos con las siguientes características:')
        
        backup_details = doc.add_paragraph()
        backup_details.add_run('• ').bold = True
        backup_details.add_run('Respaldos automáticos programados\n')
        backup_details.add_run('• ').bold = True
        backup_details.add_run('Políticas de retención configurables\n')
        backup_details.add_run('• ').bold = True
        backup_details.add_run('Cifrado de respaldos\n')
        backup_details.add_run('• ').bold = True
        backup_details.add_run('Restauración point-in-time\n')
        backup_details.add_run('• ').bold = True
        backup_details.add_run('Monitoreo y alertas de respaldos\n')
    
    else:
        # Generic service handling
        doc.add_heading(f'Configuración de {service_type.upper()}', level=1)
        doc.add_paragraph(f'La solución incluye la configuración de {service_type.upper()} con características optimizadas para el proyecto {project_name}.')
        
        generic_details = doc.add_paragraph()
        generic_details.add_run('• ').bold = True
        generic_details.add_run('Configuración siguiendo mejores prácticas de AWS\n')
        generic_details.add_run('• ').bold = True
        generic_details.add_run('Seguridad y cifrado implementados\n')
        generic_details.add_run('• ').bold = True
        generic_details.add_run('Monitoreo y alertas configurados\n')
        generic_details.add_run('• ').bold = True
        generic_details.add_run('Escalabilidad y alta disponibilidad\n')
    
    # AWS Services
    aws_services = project_info.get('aws_services', [])
    if aws_services:
        doc.add_heading('Servicios AWS Utilizados', level=1)
        for service in aws_services:
            p = doc.add_paragraph(service, style='List Bullet')
    
    # Requirements
    requirements = project_info.get('requirements', [])
    if requirements:
        doc.add_heading('Requerimientos Cumplidos', level=1)
        for req in requirements:
            p = doc.add_paragraph(req, style='List Bullet')
    
    # Architecture Overview
    doc.add_heading('Visión General de la Arquitectura', level=1)
    if service_type == 'vpc':
        doc.add_paragraph(
            'La arquitectura de red propuesta utiliza las mejores prácticas de AWS para '
            'garantizar seguridad, escalabilidad y alta disponibilidad. La VPC está '
            'diseñada con subredes públicas y privadas distribuidas en múltiples zonas '
            'de disponibilidad para máxima resiliencia.'
        )
    elif service_type == 'ec2':
        doc.add_paragraph(
            'La arquitectura de cómputo propuesta utiliza instancias EC2 optimizadas '
            'para el caso de uso específico, con configuraciones de seguridad robustas '
            'y monitoreo continuo para garantizar el rendimiento óptimo.'
        )
    elif service_type == 'rds':
        doc.add_paragraph(
            'La arquitectura de base de datos propuesta utiliza Amazon RDS para '
            'proporcionar una solución de base de datos completamente administrada '
            'con alta disponibilidad, respaldos automáticos y seguridad avanzada.'
        )
    else:
        doc.add_paragraph(
            'La solución propuesta utiliza una arquitectura moderna y escalable que aprovecha '
            'los servicios nativos de AWS para garantizar alta disponibilidad, seguridad y '
            'rendimiento óptimo.'
        )
    
    # Implementation Timeline
    doc.add_heading('Cronograma de Implementación', level=1)
    doc.add_paragraph('La implementación se realizará en las siguientes fases:')
    
    timeline = doc.add_paragraph()
    timeline.add_run('Fase 1: ').bold = True
    timeline.add_run('Configuración inicial y preparación del entorno\n')
    timeline.add_run('Fase 2: ').bold = True
    timeline.add_run('Despliegue de la infraestructura principal\n')
    timeline.add_run('Fase 3: ').bold = True
    timeline.add_run('Configuración de seguridad y monitoreo\n')
    timeline.add_run('Fase 4: ').bold = True
    timeline.add_run('Pruebas y validación\n')
    timeline.add_run('Fase 5: ').bold = True
    timeline.add_run('Puesta en producción y documentación\n')
    
    # Next Steps
    doc.add_heading('Próximos Pasos', level=1)
    doc.add_paragraph(
        'Una vez aprobada esta propuesta, procederemos con la implementación '
        'siguiendo el cronograma establecido y manteniendo comunicación constante '
        'para asegurar el éxito del proyecto.'
    )
    
    # Save to bytes
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return doc_io.getvalue()

def generate_technical_document(project_info: Dict[str, Any]) -> bytes:
    """Generate detailed technical document"""
    doc = Document()
    
    # Get project details
    project_name = project_info.get('name', 'Proyecto AWS')
    service_type = project_info.get('service_type', 'general')
    
    # Title
    doc.add_heading(f'Documento Técnico - {project_name}', 0)
    
    # Technical Overview
    doc.add_heading('Resumen Técnico', level=1)
    doc.add_paragraph(
        f'Este documento presenta los detalles técnicos de implementación '
        f'para el proyecto "{project_name}". Incluye especificaciones técnicas, '
        f'configuraciones y procedimientos de despliegue.'
    )
    
    # Technical Specifications
    doc.add_heading('Especificaciones Técnicas', level=1)
    
    if service_type == 'ec2':
        instance_type = project_info.get('instance_type', 't2.micro')
        doc.add_paragraph(f'Tipo de instancia: {instance_type}')
        doc.add_paragraph('Sistema operativo: Amazon Linux 2')
        doc.add_paragraph('Configuración de red: VPC por defecto')
        doc.add_paragraph('Almacenamiento: EBS gp3')
    elif service_type == 'vpc':
        doc.add_paragraph('Configuración de red personalizada')
        doc.add_paragraph('Subredes públicas y privadas')
        doc.add_paragraph('Internet Gateway y NAT Gateway')
    
    # Implementation Steps
    doc.add_heading('Pasos de Implementación', level=1)
    doc.add_paragraph('1. Preparación del entorno')
    doc.add_paragraph('2. Configuración de recursos')
    doc.add_paragraph('3. Despliegue de la solución')
    doc.add_paragraph('4. Pruebas y validación')
    doc.add_paragraph('5. Documentación y entrega')
    
    # Save to bytes
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return doc_io.getvalue()

def generate_technical_document(project_info: Dict[str, Any]) -> bytes:
    """Generate detailed technical document"""
    doc = Document()
    
    # Title
    doc.add_heading(f'Documento Tecnico - {project_info.get("name", "Proyecto AWS")}', 0)
    
    # Technical Architecture
    doc.add_heading('Arquitectura Tecnica Detallada', level=1)
    doc.add_paragraph(
        'Esta seccion describe en detalle la arquitectura tecnica propuesta, '
        'incluyendo todos los componentes, servicios y configuraciones necesarias.'
    )
    
    # Security Considerations
    doc.add_heading('Consideraciones de Seguridad', level=1)
    doc.add_paragraph(
        'La solucion implementa las mejores practicas de seguridad de AWS, '
        'incluyendo cifrado en transito y en reposo, control de acceso granular '
        'y monitoreo continuo.'
    )
    
    # Scalability and Performance
    doc.add_heading('Escalabilidad y Rendimiento', level=1)
    doc.add_paragraph(
        'La arquitectura esta diseñada para escalar automaticamente segun la demanda, '
        'utilizando servicios nativos de AWS para garantizar rendimiento optimo.'
    )
    
    # Monitoring and Logging
    doc.add_heading('Monitoreo y Logging', level=1)
    doc.add_paragraph(
        'Se implementaran soluciones completas de monitoreo utilizando CloudWatch, '
        'X-Ray y otros servicios de observabilidad de AWS.'
    )
    
    # Save to bytes
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return doc_io.getvalue()
