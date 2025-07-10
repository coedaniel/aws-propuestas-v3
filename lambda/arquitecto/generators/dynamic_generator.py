"""
Dynamic document generator that creates content based on AI-extracted requirements
No hardcoded services - completely dynamic and intelligent
"""
import io
import yaml
from typing import Dict, Any, List
from docx import Document

def generate_dynamic_word_document(project_info: Dict[str, Any], ai_analysis: str) -> bytes:
    """
    Generate Word document dynamically based on AI analysis and project info
    
    Args:
        project_info: Basic project information
        ai_analysis: AI's analysis of the conversation and requirements
    
    Returns:
        bytes: Word document content
    """
    doc = Document()
    
    project_name = project_info.get('name', 'Proyecto AWS')
    
    # Title
    doc.add_heading(f'Propuesta Ejecutiva - {project_name}', 0)
    
    # Executive Summary
    doc.add_heading('Resumen Ejecutivo', level=1)
    doc.add_paragraph(
        f'Este documento presenta la propuesta técnica y comercial para el proyecto '
        f'"{project_name}". La solución ha sido diseñada específicamente basándose '
        f'en los requerimientos identificados durante la consultoría.'
    )
    
    # AI Analysis Section
    doc.add_heading('Análisis de Requerimientos', level=1)
    doc.add_paragraph(
        'Basándose en la conversación mantenida, se han identificado los siguientes '
        'requerimientos y especificaciones técnicas:'
    )
    
    # Parse AI analysis for specific requirements
    analysis_lines = ai_analysis.split('\n')
    for line in analysis_lines:
        if line.strip() and not line.startswith('#'):
            doc.add_paragraph(f'• {line.strip()}', style='List Bullet')
    
    # Technical Solution
    doc.add_heading('Solución Técnica Propuesta', level=1)
    
    # Extract services mentioned in AI analysis
    services_mentioned = extract_services_from_analysis(ai_analysis)
    if services_mentioned:
        doc.add_paragraph('La solución incluye los siguientes servicios de AWS:')
        for service in services_mentioned:
            doc.add_paragraph(f'• {service}', style='List Bullet')
    
    # Implementation approach
    doc.add_heading('Enfoque de Implementación', level=1)
    doc.add_paragraph(
        'La implementación seguirá las mejores prácticas de AWS y se realizará '
        'en fases para minimizar riesgos y asegurar una transición suave.'
    )
    
    # Add phases based on complexity
    phases = generate_implementation_phases(ai_analysis, services_mentioned)
    for i, phase in enumerate(phases, 1):
        doc.add_paragraph(f'{i}. {phase}', style='List Number')
    
    # Benefits
    doc.add_heading('Beneficios de la Solución', level=1)
    benefits = extract_benefits_from_analysis(ai_analysis)
    for benefit in benefits:
        doc.add_paragraph(f'• {benefit}', style='List Bullet')
    
    # Next Steps
    doc.add_heading('Próximos Pasos', level=1)
    doc.add_paragraph('1. Revisión y aprobación de la propuesta')
    doc.add_paragraph('2. Definición del cronograma detallado')
    doc.add_paragraph('3. Inicio de la fase de implementación')
    doc.add_paragraph('4. Configuración del entorno')
    doc.add_paragraph('5. Despliegue y pruebas')
    
    # Save to bytes
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return doc_io.getvalue()

def generate_dynamic_cloudformation(project_info: Dict[str, Any], ai_analysis: str) -> str:
    """
    Generate CloudFormation template dynamically based on AI analysis
    
    Args:
        project_info: Basic project information
        ai_analysis: AI's analysis of requirements
    
    Returns:
        str: CloudFormation template in YAML format
    """
    project_name = project_info.get('name', 'aws-project').lower().replace(' ', '-')
    services = extract_services_from_analysis(ai_analysis)
    
    template = {
        'AWSTemplateFormatVersion': '2010-09-09',
        'Description': f'CloudFormation template for {project_name} - Generated dynamically from AI analysis',
        'Parameters': generate_dynamic_parameters(services, ai_analysis),
        'Resources': generate_dynamic_resources(services, project_name, ai_analysis),
        'Outputs': generate_dynamic_outputs(services, project_name)
    }
    
    return yaml.dump(template, default_flow_style=False, allow_unicode=True)

def generate_dynamic_activities_csv(project_info: Dict[str, Any], ai_analysis: str) -> bytes:
    """Generate activities CSV dynamically based on AI analysis"""
    
    import csv
    import io
    
    output = io.StringIO()
    writer = csv.writer(output)
    
    # Header
    writer.writerow(['Fase', 'Actividad', 'Descripcion', 'Duracion_Estimada', 'Responsable'])
    
    # Generate activities based on services identified
    services = extract_services_from_analysis(ai_analysis)
    activities = generate_activities_from_services(services, ai_analysis)
    
    for activity in activities:
        writer.writerow(activity)
    
    # Convert to bytes
    csv_content = output.getvalue()
    output.close()
    
    return csv_content.encode('utf-8')

def extract_services_from_analysis(ai_analysis: str) -> List[str]:
    """Extract AWS services mentioned in AI analysis"""
    services = []
    analysis_lower = ai_analysis.lower()
    
    # Common AWS services
    service_keywords = {
        'ec2': 'Amazon EC2',
        's3': 'Amazon S3',
        'rds': 'Amazon RDS',
        'lambda': 'AWS Lambda',
        'vpc': 'Amazon VPC',
        'efs': 'Amazon EFS',
        'cloudfront': 'Amazon CloudFront',
        'elb': 'Elastic Load Balancer',
        'ses': 'Amazon SES',
        'sns': 'Amazon SNS',
        'sqs': 'Amazon SQS',
        'dynamodb': 'Amazon DynamoDB',
        'cloudwatch': 'Amazon CloudWatch',
        'iam': 'AWS IAM',
        'route53': 'Amazon Route 53',
        'api gateway': 'Amazon API Gateway',
        'cognito': 'Amazon Cognito'
    }
    
    for keyword, service_name in service_keywords.items():
        if keyword in analysis_lower:
            services.append(service_name)
    
    return list(set(services))  # Remove duplicates

def generate_implementation_phases(ai_analysis: str, services: List[str]) -> List[str]:
    """Generate implementation phases based on services and complexity"""
    phases = [
        'Análisis y planificación detallada',
        'Configuración del entorno base de AWS'
    ]
    
    if services:
        phases.append(f'Implementación de servicios principales: {", ".join(services[:3])}')
        if len(services) > 3:
            phases.append(f'Configuración de servicios adicionales: {", ".join(services[3:])}')
    
    phases.extend([
        'Pruebas de integración y validación',
        'Documentación y transferencia de conocimiento',
        'Go-live y soporte post-implementación'
    ])
    
    return phases

def extract_benefits_from_analysis(ai_analysis: str) -> List[str]:
    """Extract benefits based on AI analysis"""
    benefits = [
        'Escalabilidad automática según demanda',
        'Reducción de costos operativos',
        'Mayor seguridad y compliance',
        'Disponibilidad y confiabilidad mejoradas'
    ]
    
    analysis_lower = ai_analysis.lower()
    
    if 'backup' in analysis_lower or 'respaldo' in analysis_lower:
        benefits.append('Respaldos automáticos y recuperación ante desastres')
    
    if 'monitor' in analysis_lower:
        benefits.append('Monitoreo proactivo y alertas en tiempo real')
    
    if 'performance' in analysis_lower or 'rendimiento' in analysis_lower:
        benefits.append('Optimización de rendimiento y latencia')
    
    return benefits

def generate_dynamic_parameters(services: List[str], ai_analysis: str) -> Dict[str, Any]:
    """Generate CloudFormation parameters based on services"""
    parameters = {
        'Environment': {
            'Type': 'String',
            'Default': 'prod',
            'AllowedValues': ['dev', 'staging', 'prod'],
            'Description': 'Environment name'
        }
    }
    
    # Add service-specific parameters dynamically
    for service in services:
        if 'EC2' in service:
            parameters['InstanceType'] = {
                'Type': 'String',
                'Default': 't3.micro',
                'Description': 'EC2 instance type'
            }
        elif 'S3' in service:
            parameters['BucketName'] = {
                'Type': 'String',
                'Description': 'S3 bucket name'
            }
        elif 'RDS' in service:
            parameters['DBInstanceClass'] = {
                'Type': 'String',
                'Default': 'db.t3.micro',
                'Description': 'RDS instance class'
            }
    
    return parameters

def generate_dynamic_resources(services: List[str], project_name: str, ai_analysis: str) -> Dict[str, Any]:
    """Generate CloudFormation resources based on services"""
    resources = {}
    
    # Generate resources dynamically based on identified services
    for service in services:
        if 'EC2' in service:
            resources['EC2Instance'] = {
                'Type': 'AWS::EC2::Instance',
                'Properties': {
                    'InstanceType': {'Ref': 'InstanceType'},
                    'ImageId': 'ami-0c02fb55956c7d316',  # Amazon Linux 2
                    'Tags': [
                        {'Key': 'Name', 'Value': f'{project_name}-ec2'},
                        {'Key': 'Project', 'Value': project_name}
                    ]
                }
            }
        elif 'S3' in service:
            resources['S3Bucket'] = {
                'Type': 'AWS::S3::Bucket',
                'Properties': {
                    'BucketName': {'Ref': 'BucketName'},
                    'Tags': [
                        {'Key': 'Name', 'Value': f'{project_name}-s3'},
                        {'Key': 'Project', 'Value': project_name}
                    ]
                }
            }
        elif 'RDS' in service:
            resources['RDSInstance'] = {
                'Type': 'AWS::RDS::DBInstance',
                'Properties': {
                    'DBInstanceClass': {'Ref': 'DBInstanceClass'},
                    'Engine': 'mysql',
                    'MasterUsername': 'admin',
                    'MasterUserPassword': 'changeme123',
                    'AllocatedStorage': '20',
                    'Tags': [
                        {'Key': 'Name', 'Value': f'{project_name}-rds'},
                        {'Key': 'Project', 'Value': project_name}
                    ]
                }
            }
    
    # If no specific resources, add a generic one
    if not resources:
        resources['GenericResource'] = {
            'Type': 'AWS::CloudFormation::WaitConditionHandle',
            'Properties': {}
        }
    
    return resources

def generate_dynamic_outputs(services: List[str], project_name: str) -> Dict[str, Any]:
    """Generate CloudFormation outputs based on services"""
    outputs = {}
    
    for service in services:
        if 'EC2' in service:
            outputs['EC2InstanceId'] = {
                'Description': 'EC2 Instance ID',
                'Value': {'Ref': 'EC2Instance'}
            }
        elif 'S3' in service:
            outputs['S3BucketName'] = {
                'Description': 'S3 Bucket Name',
                'Value': {'Ref': 'S3Bucket'}
            }
        elif 'RDS' in service:
            outputs['RDSEndpoint'] = {
                'Description': 'RDS Endpoint',
                'Value': {'Fn::GetAtt': ['RDSInstance', 'Endpoint.Address']}
            }
    
    if not outputs:
        outputs['ProjectName'] = {
            'Description': 'Project Name',
            'Value': project_name
        }
    
    return outputs

def generate_activities_from_services(services: List[str], ai_analysis: str) -> List[List[str]]:
    """Generate implementation activities based on services"""
    activities = [
        ['Planificacion', 'Analisis de requerimientos', 'Revision detallada de necesidades del proyecto', '2 dias', 'Arquitecto AWS'],
        ['Planificacion', 'Diseño de arquitectura', 'Diseño de la solucion tecnica completa', '3 dias', 'Arquitecto AWS']
    ]
    
    for service in services:
        if 'EC2' in service:
            activities.extend([
                ['Implementacion', 'Configuracion EC2', 'Creacion y configuracion de instancias EC2', '1 dia', 'DevOps Engineer'],
                ['Implementacion', 'Configuracion seguridad EC2', 'Configuracion de security groups y accesos', '0.5 dias', 'Security Engineer']
            ])
        elif 'S3' in service:
            activities.extend([
                ['Implementacion', 'Configuracion S3', 'Creacion y configuracion de buckets S3', '0.5 dias', 'DevOps Engineer'],
                ['Implementacion', 'Politicas S3', 'Configuracion de politicas de acceso y cifrado', '0.5 dias', 'Security Engineer']
            ])
        elif 'RDS' in service:
            activities.extend([
                ['Implementacion', 'Configuracion RDS', 'Creacion y configuracion de base de datos RDS', '1 dia', 'Database Administrator'],
                ['Implementacion', 'Backup RDS', 'Configuracion de respaldos automaticos', '0.5 dias', 'Database Administrator']
            ])
    
    activities.extend([
        ['Pruebas', 'Pruebas de integracion', 'Validacion de todos los componentes', '2 dias', 'QA Engineer'],
        ['Entrega', 'Documentacion', 'Creacion de documentacion tecnica', '1 dia', 'Technical Writer'],
        ['Entrega', 'Transferencia conocimiento', 'Capacitacion al equipo cliente', '1 dia', 'Arquitecto AWS']
    ])
    
    return activities
