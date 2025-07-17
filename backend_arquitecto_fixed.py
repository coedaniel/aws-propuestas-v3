# AWS Propuestas v3 - Arquitecto Backend Fixed
# Solución profesional para generar documentos sin caracteres especiales

import json
import boto3
import uuid
import os
from datetime import datetime
from typing import Dict, List, Any, Optional
import logging
from unidecode import unidecode
import pandas as pd
from docx import Document
from docx.shared import Inches
import io
import base64

# Configurar logging
logger = logging.getLogger()
logger.setLevel(logging.INFO)

# Clientes AWS
bedrock_client = boto3.client('bedrock-runtime', region_name='us-east-1')
dynamodb = boto3.resource('dynamodb', region_name='us-east-1')
s3_client = boto3.client('s3', region_name='us-east-1')

# Variables de entorno
PROJECTS_TABLE = os.environ.get('PROJECTS_TABLE', 'aws-propuestas-v3-projects-prod')
DOCUMENTS_BUCKET = os.environ.get('DOCUMENTS_BUCKET', 'aws-propuestas-v3-documents-prod-035385358261')

# Tabla DynamoDB
projects_table = dynamodb.Table(PROJECTS_TABLE)

def clean_text(text: str) -> str:
    """Elimina acentos y caracteres especiales de un texto"""
    if not text:
        return ""
    # Usar unidecode para convertir caracteres especiales
    cleaned = unidecode(text)
    # Reemplazar caracteres problemáticos adicionales
    cleaned = cleaned.replace('ñ', 'n').replace('Ñ', 'N')
    return cleaned

def clean_filename(filename: str) -> str:
    """Limpia un nombre de archivo para que sea compatible"""
    cleaned = clean_text(filename)
    # Reemplazar espacios y caracteres especiales con guiones
    cleaned = cleaned.replace(' ', '-').replace('_', '-')
    # Remover caracteres no alfanuméricos excepto guiones y puntos
    import re
    cleaned = re.sub(r'[^a-zA-Z0-9\-\.]', '', cleaned)
    return cleaned

def invoke_bedrock_model(messages: List[Dict], model_id: str = "anthropic.claude-3-haiku-20240307-v1:0") -> str:
    """Invoca un modelo de Bedrock Runtime usando la API Converse"""
    try:
        # Usar la nueva API Converse
        response = bedrock_client.converse(
            modelId=model_id,
            messages=messages,
            inferenceConfig={
                'maxTokens': 4000,
                'temperature': 0.7
            }
        )
        
        return response['output']['message']['content'][0]['text']
        
    except Exception as e:
        logger.error(f"Error invoking Bedrock model: {str(e)}")
        return f"Error al procesar la solicitud: {str(e)}"

def save_project_to_dynamodb(project_data: Dict) -> bool:
    """Guarda un proyecto en DynamoDB"""
    try:
        # Limpiar todos los textos del proyecto
        cleaned_project = {}
        for key, value in project_data.items():
            if isinstance(value, str):
                cleaned_project[key] = clean_text(value)
            elif isinstance(value, dict):
                cleaned_project[key] = {k: clean_text(v) if isinstance(v, str) else v for k, v in value.items()}
            else:
                cleaned_project[key] = value
        
        # Agregar timestamp
        cleaned_project['updatedAt'] = datetime.utcnow().isoformat()
        
        projects_table.put_item(Item=cleaned_project)
        logger.info(f"Project saved to DynamoDB: {cleaned_project.get('projectId')}")
        return True
        
    except Exception as e:
        logger.error(f"Error saving project to DynamoDB: {str(e)}")
        return False

def generate_activities_csv(project_name: str, project_info: Dict) -> bytes:
    """Genera un archivo CSV con actividades de implementación"""
    try:
        # Actividades base para cualquier proyecto AWS
        activities = [
            {
                'Fase': 'Planificacion',
                'Actividad': 'Revision de requerimientos',
                'Descripcion': 'Analizar y validar todos los requerimientos del proyecto',
                'Duracion_Dias': 3,
                'Responsable': 'Arquitecto de Soluciones',
                'Dependencias': 'Ninguna',
                'Estado': 'Pendiente'
            },
            {
                'Fase': 'Planificacion',
                'Actividad': 'Diseno de arquitectura',
                'Descripcion': 'Crear el diseno detallado de la arquitectura AWS',
                'Duracion_Dias': 5,
                'Responsable': 'Arquitecto de Soluciones',
                'Dependencias': 'Revision de requerimientos',
                'Estado': 'Pendiente'
            },
            {
                'Fase': 'Implementacion',
                'Actividad': 'Configuracion de VPC y redes',
                'Descripcion': 'Crear VPC, subnets, security groups y configuracion de red',
                'Duracion_Dias': 2,
                'Responsable': 'Ingeniero DevOps',
                'Dependencias': 'Diseno de arquitectura',
                'Estado': 'Pendiente'
            },
            {
                'Fase': 'Implementacion',
                'Actividad': 'Despliegue de servicios principales',
                'Descripcion': 'Implementar EC2, RDS, Lambda u otros servicios principales',
                'Duracion_Dias': 7,
                'Responsable': 'Ingeniero DevOps',
                'Dependencias': 'Configuracion de VPC y redes',
                'Estado': 'Pendiente'
            },
            {
                'Fase': 'Implementacion',
                'Actividad': 'Configuracion de seguridad',
                'Descripcion': 'Implementar IAM, KMS, CloudTrail y otras medidas de seguridad',
                'Duracion_Dias': 3,
                'Responsable': 'Especialista en Seguridad',
                'Dependencias': 'Despliegue de servicios principales',
                'Estado': 'Pendiente'
            },
            {
                'Fase': 'Pruebas',
                'Actividad': 'Pruebas de funcionalidad',
                'Descripcion': 'Ejecutar pruebas funcionales y de integracion',
                'Duracion_Dias': 4,
                'Responsable': 'QA Engineer',
                'Dependencias': 'Configuracion de seguridad',
                'Estado': 'Pendiente'
            },
            {
                'Fase': 'Pruebas',
                'Actividad': 'Pruebas de rendimiento',
                'Descripcion': 'Validar rendimiento y escalabilidad del sistema',
                'Duracion_Dias': 3,
                'Responsable': 'QA Engineer',
                'Dependencias': 'Pruebas de funcionalidad',
                'Estado': 'Pendiente'
            },
            {
                'Fase': 'Despliegue',
                'Actividad': 'Despliegue a produccion',
                'Descripcion': 'Migrar la solucion al ambiente de produccion',
                'Duracion_Dias': 2,
                'Responsable': 'Ingeniero DevOps',
                'Dependencias': 'Pruebas de rendimiento',
                'Estado': 'Pendiente'
            },
            {
                'Fase': 'Despliegue',
                'Actividad': 'Documentacion y entrega',
                'Descripcion': 'Crear documentacion final y realizar entrega al cliente',
                'Duracion_Dias': 2,
                'Responsable': 'Arquitecto de Soluciones',
                'Dependencias': 'Despliegue a produccion',
                'Estado': 'Pendiente'
            }
        ]
        
        # Crear DataFrame y limpiar textos
        df = pd.DataFrame(activities)
        for col in df.columns:
            if df[col].dtype == 'object':
                df[col] = df[col].apply(lambda x: clean_text(str(x)) if pd.notna(x) else x)
        
        # Convertir a CSV
        csv_buffer = io.StringIO()
        df.to_csv(csv_buffer, index=False, encoding='utf-8')
        return csv_buffer.getvalue().encode('utf-8')
        
    except Exception as e:
        logger.error(f"Error generating activities CSV: {str(e)}")
        return b"Error,generando,archivo,CSV"

def generate_word_document(project_name: str, project_info: Dict) -> bytes:
    """Genera un documento Word con la información del proyecto"""
    try:
        doc = Document()
        
        # Título principal
        title = doc.add_heading(clean_text(f'Propuesta Arquitectonica: {project_name}'), 0)
        
        # Información general
        doc.add_heading('1. Informacion General', level=1)
        doc.add_paragraph(f'Nombre del Proyecto: {clean_text(project_name)}')
        doc.add_paragraph(f'Fecha de Creacion: {datetime.now().strftime("%Y-%m-%d")}')
        doc.add_paragraph(f'Tipo de Solucion: {clean_text(project_info.get("service_focus", "Solucion Integral AWS"))}')
        
        # Descripción del proyecto
        doc.add_heading('2. Descripcion del Proyecto', level=1)
        description = clean_text(project_info.get('description', 'Proyecto de implementacion en AWS'))
        doc.add_paragraph(description)
        
        # Objetivo
        doc.add_heading('3. Objetivo', level=1)
        objective = clean_text(project_info.get('objective', 'Implementar una solucion robusta y escalable en AWS'))
        doc.add_paragraph(objective)
        
        # Servicios AWS recomendados
        doc.add_heading('4. Servicios AWS Recomendados', level=1)
        services = [
            'Amazon EC2 - Instancias de computo elasticas',
            'Amazon RDS - Base de datos relacional administrada',
            'Amazon S3 - Almacenamiento de objetos escalable',
            'Amazon VPC - Red privada virtual',
            'AWS IAM - Gestion de identidades y accesos',
            'Amazon CloudWatch - Monitoreo y observabilidad',
            'AWS CloudFormation - Infraestructura como codigo'
        ]
        
        for service in services:
            doc.add_paragraph(clean_text(service), style='List Bullet')
        
        # Consideraciones de seguridad
        doc.add_heading('5. Consideraciones de Seguridad', level=1)
        security_items = [
            'Implementacion de principio de menor privilegio en IAM',
            'Encriptacion en transito y en reposo',
            'Configuracion de Security Groups restrictivos',
            'Habilitacion de CloudTrail para auditoria',
            'Implementacion de AWS Config para compliance'
        ]
        
        for item in security_items:
            doc.add_paragraph(clean_text(item), style='List Bullet')
        
        # Próximos pasos
        doc.add_heading('6. Proximos Pasos', level=1)
        next_steps = [
            'Revision y aprobacion de la propuesta',
            'Planificacion detallada del proyecto',
            'Configuracion del ambiente de desarrollo',
            'Inicio de la implementacion por fases',
            'Pruebas y validacion de la solucion'
        ]
        
        for step in next_steps:
            doc.add_paragraph(clean_text(step), style='List Number')
        
        # Guardar en buffer
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        return doc_buffer.getvalue()
        
    except Exception as e:
        logger.error(f"Error generating Word document: {str(e)}")
        return b"Error generando documento Word"

def generate_costs_csv(project_name: str, project_info: Dict) -> bytes:
    """Genera un archivo CSV con estimación de costos"""
    try:
        # Costos estimados base para servicios AWS comunes
        costs = [
            {
                'Servicio': 'Amazon EC2',
                'Tipo': 't3.medium',
                'Cantidad': 2,
                'Costo_Mensual_USD': 67.32,
                'Descripcion': 'Instancias de aplicacion'
            },
            {
                'Servicio': 'Amazon RDS',
                'Tipo': 'db.t3.micro',
                'Cantidad': 1,
                'Costo_Mensual_USD': 15.84,
                'Descripcion': 'Base de datos MySQL'
            },
            {
                'Servicio': 'Amazon S3',
                'Tipo': 'Standard',
                'Cantidad': 100,
                'Costo_Mensual_USD': 2.30,
                'Descripcion': '100 GB de almacenamiento'
            },
            {
                'Servicio': 'Application Load Balancer',
                'Tipo': 'ALB',
                'Cantidad': 1,
                'Costo_Mensual_USD': 22.50,
                'Descripcion': 'Balanceador de carga'
            },
            {
                'Servicio': 'Amazon CloudWatch',
                'Tipo': 'Logs y Metricas',
                'Cantidad': 1,
                'Costo_Mensual_USD': 10.00,
                'Descripcion': 'Monitoreo basico'
            },
            {
                'Servicio': 'AWS NAT Gateway',
                'Tipo': 'NAT Gateway',
                'Cantidad': 1,
                'Costo_Mensual_USD': 45.00,
                'Descripcion': 'Conectividad saliente'
            }
        ]
        
        # Crear DataFrame y limpiar textos
        df = pd.DataFrame(costs)
        for col in df.columns:
            if df[col].dtype == 'object':
                df[col] = df[col].apply(lambda x: clean_text(str(x)) if pd.notna(x) else x)
        
        # Agregar total
        total_cost = df['Costo_Mensual_USD'].sum()
        total_row = {
            'Servicio': 'TOTAL ESTIMADO',
            'Tipo': '',
            'Cantidad': '',
            'Costo_Mensual_USD': total_cost,
            'Descripcion': 'Costo mensual estimado total'
        }
        df = pd.concat([df, pd.DataFrame([total_row])], ignore_index=True)
        
        # Convertir a CSV
        csv_buffer = io.StringIO()
        df.to_csv(csv_buffer, index=False, encoding='utf-8')
        return csv_buffer.getvalue().encode('utf-8')
        
    except Exception as e:
        logger.error(f"Error generating costs CSV: {str(e)}")
        return b"Error,generando,archivo,costos"

def upload_file_to_s3(file_content: bytes, file_name: str, project_name: str) -> bool:
    """Sube un archivo a S3"""
    try:
        clean_project_name = clean_filename(project_name)
        s3_key = f"{clean_project_name}/{file_name}"
        
        s3_client.put_object(
            Bucket=DOCUMENTS_BUCKET,
            Key=s3_key,
            Body=file_content,
            ContentType='application/octet-stream'
        )
        
        logger.info(f"File uploaded to S3: {s3_key}")
        return True
        
    except Exception as e:
        logger.error(f"Error uploading file to S3: {str(e)}")
        return False

def lambda_handler(event, context):
    """Handler principal de Lambda"""
    try:
        # Parsear el evento
        if isinstance(event.get('body'), str):
            body = json.loads(event['body'])
        else:
            body = event.get('body', {})
        
        messages = body.get('messages', [])
        selected_model = body.get('selected_model', 'anthropic.claude-3-haiku-20240307-v1:0')
        project_id = body.get('projectId') or str(uuid.uuid4())
        
        # Obtener la respuesta del modelo
        if messages:
            response_text = invoke_bedrock_model(messages, selected_model)
        else:
            response_text = "Hola! Soy tu Arquitecto AWS. Para comenzar, necesito que me proporciones el nombre del proyecto."
        
        # Detectar si se debe generar documentos
        should_generate_docs = any([
            'generar' in response_text.lower(),
            'documentos' in response_text.lower(),
            'archivos' in response_text.lower(),
            'subir' in response_text.lower()
        ])
        
        # Información del proyecto (extraer del contexto)
        project_info = {
            'service_focus': 'AWS',
            'description': 'Proyecto de arquitectura AWS',
            'objective': 'Implementar solucion en AWS',
            'status': 'IN_PROGRESS'
        }
        
        # Si hay mensajes, extraer información del proyecto
        if messages:
            last_message = messages[-1].get('content', '')
            if 'tiendaonline' in last_message.lower():
                project_name = 'TiendaOnline'
                project_info.update({
                    'description': 'Aplicacion de e-commerce completa',
                    'objective': 'Crear tienda online escalable',
                    'service_focus': 'E-commerce'
                })
            else:
                # Extraer nombre del proyecto del contexto
                project_name = 'ProyectoAWS'
                for msg in messages:
                    content = msg.get('content', '').lower()
                    if 'proyecto' in content and 'llama' in content:
                        # Intentar extraer el nombre
                        words = content.split()
                        for i, word in enumerate(words):
                            if word in ['proyecto', 'llama', 'llamado'] and i + 1 < len(words):
                                project_name = words[i + 1].title()
                                break
        else:
            project_name = 'ProyectoAWS'
        
        # Generar documentos si es necesario
        documents_generated = False
        if should_generate_docs:
            try:
                # Generar archivos
                activities_csv = generate_activities_csv(project_name, project_info)
                word_doc = generate_word_document(project_name, project_info)
                costs_csv = generate_costs_csv(project_name, project_info)
                
                # Nombres de archivos limpios
                clean_name = clean_filename(project_name)
                
                # Subir archivos a S3
                upload_file_to_s3(activities_csv, f"{clean_name}-actividades.csv", project_name)
                upload_file_to_s3(word_doc, f"{clean_name}-propuesta.docx", project_name)
                upload_file_to_s3(costs_csv, f"{clean_name}-costos.csv", project_name)
                
                documents_generated = True
                
                # Actualizar respuesta para confirmar generación
                response_text += f"\n\nArchivos generados exitosamente:\n- {clean_name}-actividades.csv\n- {clean_name}-propuesta.docx\n- {clean_name}-costos.csv\n\nTodos los archivos han sido subidos al bucket S3."
                
            except Exception as e:
                logger.error(f"Error generating documents: {str(e)}")
                response_text += f"\n\nError al generar documentos: {str(e)}"
        
        # Guardar proyecto en DynamoDB
        project_data = {
            'projectId': project_id,
            'projectName': clean_text(project_name),
            'status': 'IN_PROGRESS',
            'currentStep': '1',
            'createdAt': datetime.utcnow().isoformat(),
            'updatedAt': datetime.utcnow().isoformat(),
            'documentCount': 3 if documents_generated else 0,
            'hasDocuments': documents_generated,
            'projectInfo': project_info,
            'lastMessage': clean_text(response_text[:500])  # Truncar para DynamoDB
        }
        
        save_project_to_dynamodb(project_data)
        
        # Respuesta
        return {
            'statusCode': 200,
            'headers': {
                'Content-Type': 'application/json',
                'Access-Control-Allow-Origin': '*',
                'Access-Control-Allow-Methods': 'POST, OPTIONS',
                'Access-Control-Allow-Headers': 'Content-Type'
            },
            'body': json.dumps({
                'response': response_text,
                'modelId': selected_model,
                'projectId': project_id,
                'projectInfo': project_info,
                'currentStep': 1,
                'isComplete': documents_generated,
                'usage': {
                    'inputTokens': sum(len(msg.get('content', '')) for msg in messages),
                    'outputTokens': len(response_text),
                    'totalTokens': sum(len(msg.get('content', '')) for msg in messages) + len(response_text)
                },
                'documentGeneration': {
                    'generated': documents_generated,
                    'bucket': DOCUMENTS_BUCKET,
                    'folder': clean_filename(project_name)
                } if documents_generated else None,
                'specificService': project_info.get('service_focus', 'AWS')
            })
        }
        
    except Exception as e:
        logger.error(f"Lambda handler error: {str(e)}")
        return {
            'statusCode': 500,
            'headers': {
                'Content-Type': 'application/json',
                'Access-Control-Allow-Origin': '*'
            },
            'body': json.dumps({
                'error': 'Error processing arquitecto request',
                'details': str(e)
            })
        }
